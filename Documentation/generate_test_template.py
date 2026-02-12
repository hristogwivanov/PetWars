from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title
title = doc.add_heading('Тестване на алгоритмите', level=2)

intro = doc.add_paragraph(
    'За сравнение на трите алгоритъма са проведени тестове с три различни разстояния: '
    'кратко, средно и дълго. За всеки тест са измерени времето за изпълнение и броят '
    'на посетените върхове.'
)

distances = ['Кратко разстояние', 'Средно разстояние', 'Дълго разстояние']
algorithms = ['Дийкстра', 'A*', 'Белман-Форд']
metrics = ['Време (ms)', 'Посетени върхове', 'Дължина на пътя']

for dist_name in distances:
    doc.add_heading(dist_name, level=3)
    
    # Placeholder for screenshot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Тук поставете скрийншот]')
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Description placeholders
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Начална позиция: (__, __) → Целева позиция: (__, __)')
    run2.font.color.rgb = RGBColor(100, 100, 100)

    # Table: algorithms as rows, metrics as columns
    table = doc.add_table(rows=4, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Алгоритъм'] + metrics
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    # Algorithm rows
    for row_idx, algo in enumerate(algorithms, start=1):
        table.rows[row_idx].cells[0].text = algo
        for col_idx in range(1, 4):
            table.rows[row_idx].cells[col_idx].text = ''

    doc.add_paragraph()  # spacing

# Summary comparison table
doc.add_heading('Обобщена таблица за сравнение', level=3)

summary = doc.add_table(rows=4, cols=7, style='Table Grid')
summary.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
sum_headers = ['Алгоритъм', 
               'Кратко\nВреме (ms)', 'Кратко\nВърхове',
               'Средно\nВреме (ms)', 'Средно\nВърхове', 
               'Дълго\nВреме (ms)', 'Дълго\nВърхове']
for i, h in enumerate(sum_headers):
    cell = summary.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

for row_idx, algo in enumerate(algorithms, start=1):
    summary.rows[row_idx].cells[0].text = algo

doc.add_paragraph()

# Conclusion placeholder
doc.add_heading('Анализ на резултатите', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '[Тук напишете кратък анализ на резултатите: кой алгоритъм е най-бърз, '
    'кой посещава най-малко върхове, как се различават при различни разстояния.]'
)
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

import os
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Test_Results_Template.docx')
doc.save(out_path)
print("Template saved to Documentation/Test_Results_Template.docx")
