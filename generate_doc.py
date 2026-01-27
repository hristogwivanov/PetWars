from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set paragraph formatting - JUSTIFIED alignment
paragraph_format = style.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Двустранно подравняване
paragraph_format.space_after = Pt(6)
paragraph_format.space_before = Pt(6)
paragraph_format.line_spacing = 1.5
paragraph_format.first_line_indent = Cm(1.25)

# Add footer with name, course, faculty number, page number
section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add text to footer - UPDATE THESE WITH YOUR INFO
footer_para.add_run("Вашето Име | Курс: X | Факултетен номер: XXXXX | Страница ")

# Add page number field
run = footer_para.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar1)

run2 = footer_para.add_run()
instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"
run2._r.append(instrText)

run3 = footer_para.add_run()
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run3._r.append(fldChar2)

# Set footer font
for run in footer_para.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# Title
title = doc.add_heading('Курсов проект', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Алгоритми за намиране на най-кратък път в граф')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 1. Theoretical Part
doc.add_heading('1. Теоретична част', level=1)

doc.add_heading('1.1 Представяне на избраните алгоритми', level=2)

p = doc.add_paragraph(
    'В настоящия проект са имплементирани три алгоритъма за намиране на най-кратък път в граф: '
    'алгоритъмът на Дийкстра, алгоритъмът A* и алгоритъмът на Белман-Форд. '
    'Тези алгоритми са фундаментални в теорията на графите и намират широко приложение в компютърните игри, '
    'навигационните системи и мрежовите протоколи.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.2 Алгоритъм на Дийкстра', level=2)

p = doc.add_paragraph(
    'Алгоритъмът на Дийкстра е класически алгоритъм за намиране на най-кратък път от един източник '
    'до всички останали върхове в граф с неотрицателни тегла на ребрата. Разработен е от холандския '
    'учен Едсхер Дийкстра през 1956 г.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Основна идея и принцип на работа:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('• Поддържа се множество от посетени върхове и приоритетна опашка с непосетени върхове')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• За всеки връх се пази текущото най-кратко разстояние от началния връх')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• На всяка стъпка се избира непосетеният връх с най-малко разстояние')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Актуализират се разстоянията до съседите на избрания връх')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph('Сложност: O((V+E) log V), където V е броят на върховете, а E е броят на ребрата.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.3 Алгоритъм A*', level=2)

p = doc.add_paragraph(
    'Алгоритъмът A* е разширение на алгоритъма на Дийкстра, което използва евристична функция '
    'за насочване на търсенето към целта. Това го прави значително по-ефективен в практически приложения.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Основна идея и принцип на работа:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Използва функция f(n) = g(n) + h(n), където:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('  - g(n) е реалната цена от началото до текущия връх')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('  - h(n) е евристичната оценка от текущия връх до целта')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• В проекта се използва Manhattan distance като евристика')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Алгоритъмът приоритизира върхове, които са по-близо до целта')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph('Сложност: O((V+E) log V) в най-лошия случай, но практически много по-бърз от Дийкстра.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.4 Алгоритъм на Белман-Форд', level=2)

p = doc.add_paragraph(
    'Алгоритъмът на Белман-Форд намира най-кратките пътища от един източник до всички върхове, '
    'като може да работи и с отрицателни тегла на ребрата (за разлика от Дийкстра).'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Основна идея и принцип на работа:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Извършва V-1 итерации, където V е броят на върховете')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• На всяка итерация "релаксира" всички ребра в графа')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Релаксация означава проверка дали пътят през дадено ребро е по-кратък')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph('Сложност: O(V*E), което е по-бавно от Дийкстра, но по-универсално.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('1.5 Сравнение на алгоритмите', level=2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Алгоритъм'
hdr_cells[1].text = 'Сложност'
hdr_cells[2].text = 'Предимства'
hdr_cells[3].text = 'Недостатъци'

# Dijkstra row
row_cells = table.rows[1].cells
row_cells[0].text = 'Дийкстра'
row_cells[1].text = 'O((V+E) log V)'
row_cells[2].text = 'Оптимален, надежден'
row_cells[3].text = 'Изследва всички посоки'

# A* row
row_cells = table.rows[2].cells
row_cells[0].text = 'A*'
row_cells[1].text = 'O((V+E) log V)'
row_cells[2].text = 'Насочено търсене, бърз'
row_cells[3].text = 'Изисква добра евристика'

# Bellman-Ford row
row_cells = table.rows[3].cells
row_cells[0].text = 'Белман-Форд'
row_cells[1].text = 'O(V*E)'
row_cells[2].text = 'Работи с отрицателни тегла'
row_cells[3].text = 'По-бавен'

# Center the table
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 2. Example
doc.add_heading('2. Пример за работата на алгоритмите', level=1)

p = doc.add_paragraph(
    'В играта PetWars алгоритмите се използват за намиране на път от текущата позиция на героя '
    'до избрана целева клетка на картата. Картата е представена като двумерна решетка, където '
    'всяка клетка може да бъде проходима или непроходима.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Стъпки на изпълнение (на примера на Дийкстра):')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('1. Инициализация: началната позиция получава разстояние 0, всички останали - безкрайност')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('2. Добавяне на началната позиция в приоритетната опашка')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('3. Извличане на върха с най-малко разстояние от опашката')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('4. За всеки съсед: ако новият път е по-кратък, актуализиране на разстоянието')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('5. Повтаряне докато се достигне целта или опашката се изпразни')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('6. Възстановяване на пътя чрез обратно проследяване')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph(
    'В демо режима на играта визуализацията показва: оранжеви клетки за посетени върхове, '
    'жълти за върхове в опашката (frontier), червена за текущо разглеждания връх и синя за '
    'намерения път.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Placeholder for figure
doc.add_paragraph()
p = doc.add_paragraph('[Тук добавете екранна снимка от демо режима]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

p = doc.add_paragraph('Фигура 1: Визуализация на алгоритъма на Дийкстра в демо режим')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

# 3. Program Implementation
doc.add_heading('3. Програмна реализация', level=1)

p = doc.add_paragraph(
    'Проектът е реализиран на Python с използване на библиотеката Pygame за графичния интерфейс. '
    'Основните файлове са:'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('• main.py - главен файл с игровия цикъл и визуализация')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• pathfinding.py - имплементация на алгоритмите')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• gamedata.py - класове за герои, ресурси и AI')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• constants.py - константи и карти на терена')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('3.1 Имплементация на алгоритъма на Дийкстра', level=2)

code = '''def dijkstra_path(start, goal, terrain_map):
    import heapq
    
    frontier = [(0, start)]
    came_from = {start: None}
    cost_so_far = {start: 0}
    
    while frontier:
        current_cost, current = heapq.heappop(frontier)
        
        if current == goal:
            break
            
        for next_node in get_neighbors(current, terrain_map):
            new_cost = cost_so_far[current] + 1
            if next_node not in cost_so_far or new_cost < cost_so_far[next_node]:
                cost_so_far[next_node] = new_cost
                priority = new_cost
                heapq.heappush(frontier, (priority, next_node))
                came_from[next_node] = current
    
    # Reconstruct path
    path = []
    current = goal
    while current is not None:
        path.append(current)
        current = came_from.get(current)
    path.reverse()
    return path'''

p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_heading('3.2 Имплементация на A* с евристика', level=2)

code_astar = '''def heuristic(a, b):
    # Manhattan distance
    return abs(a[0] - b[0]) + abs(a[1] - b[1])

def astar_path(start, goal, terrain_map):
    import heapq
    
    frontier = [(0, start)]
    came_from = {start: None}
    cost_so_far = {start: 0}
    
    while frontier:
        current_cost, current = heapq.heappop(frontier)
        
        if current == goal:
            break
            
        for next_node in get_neighbors(current, terrain_map):
            new_cost = cost_so_far[current] + 1
            if next_node not in cost_so_far or new_cost < cost_so_far[next_node]:
                cost_so_far[next_node] = new_cost
                priority = new_cost + heuristic(next_node, goal)
                heapq.heappush(frontier, (priority, next_node))
                came_from[next_node] = current
    
    return reconstruct_path(came_from, goal)'''

p = doc.add_paragraph()
run = p.add_run(code_astar)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_heading('3.3 Имплементация на Белман-Форд', level=2)

code_bf = '''def bellman_ford_path(start, goal, terrain_map):
    # Initialize distances
    dist = {}
    came_from = {}
    
    # Get all walkable nodes
    nodes = []
    for y in range(len(terrain_map)):
        for x in range(len(terrain_map[0])):
            if terrain_map[y][x] > 0:
                nodes.append((x, y))
                dist[(x, y)] = float('inf')
    
    dist[start] = 0
    
    # Relax edges V-1 times
    for _ in range(len(nodes) - 1):
        for node in nodes:
            if dist[node] == float('inf'):
                continue
            for neighbor in get_neighbors(node, terrain_map):
                if dist[node] + 1 < dist[neighbor]:
                    dist[neighbor] = dist[node] + 1
                    came_from[neighbor] = node
    
    return reconstruct_path(came_from, start, goal)'''

p = doc.add_paragraph()
run = p.add_run(code_bf)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 4. Demonstration
doc.add_heading('4. Демонстрация на работа', level=1)

p = doc.add_paragraph('Входни данни:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Карта на терена (20x18 клетки)')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Начална позиция на героя')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Целева позиция (избрана с клик)')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph('Изходни данни:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Визуализация на процеса на търсене')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Намереният най-кратък път')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Брой посетени върхове')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Време за изпълнение (в милисекунди)')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph(
    'Демо режимът се активира с бутона "DEMO MODE" или клавиш D. '
    'В този режим алгоритъмът се визуализира стъпка по стъпка, '
    'като потребителят може да избира между трите алгоритъма.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Placeholder for figure 2
doc.add_paragraph()
p = doc.add_paragraph('[Тук добавете екранна снимка с избор на алгоритъм]')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

p = doc.add_paragraph('Фигура 2: Интерфейс за избор на алгоритъм и показване на статистика')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

# 5. Conclusion
doc.add_heading('5. Заключение', level=1)

p = doc.add_paragraph(
    'В проекта успешно са имплементирани три алгоритъма за намиране на най-кратък път: '
    'Дийкстра, A* и Белман-Форд. Визуализацията позволява да се наблюдава работата на '
    'алгоритмите в реално време и да се сравни тяхната ефективност.'
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Основни резултати:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• A* е най-бърз за насочено търсене към конкретна цел')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Дийкстра изследва по-голяма област, но гарантира оптималност')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Белман-Форд е по-бавен, но работи с по-общи случаи')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
p = doc.add_paragraph('Възможности за бъдещи разработки:')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Добавяне на различни тегла на терена (вода, планини)')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Имплементация на Jump Point Search за още по-бързо търсене')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('• Паралелизация на алгоритмите за по-големи карти')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 6. References
doc.add_heading('6. Използвана литература', level=1)

p = doc.add_paragraph('1. Cormen, T. H., et al. "Introduction to Algorithms", MIT Press, 2009')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('2. Russell, S., Norvig, P. "Artificial Intelligence: A Modern Approach", Pearson, 2020')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('3. Pygame Documentation - https://www.pygame.org/docs/')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph('4. Red Blob Games - Introduction to A* - https://www.redblobgames.com/pathfinding/a-star/')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Save document
doc.save('d:\\ico\\VFU Master Degree\\programming-and-algorithms-vfu-final-project\\PetWars\\Курсов_проект_Pathfinding.docx')
print('Document created successfully!')
print('')
print('ВАЖНО: Отворете документа и:')
print('1. Заменете "Вашето Име | Курс: X | Факултетен номер: XXXXX" в долния колонтитул')
print('2. Добавете екранни снимки на местата, маркирани с [Тук добавете...]')
